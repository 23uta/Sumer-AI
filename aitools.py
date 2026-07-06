# ==========
# libraries
# ==========

import os
import json
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import subprocess
import platform

#-----------------
#files calling :
#-----------------
from configs import client
import controltools 
from configs import excel_prompt


#===================
#making reports defs
#===================

# ════════════════════════════════════════════════════════════
#  1. إعدادات الصفحة
# ════════════════════════════════════════════════════════════

def setup_page_size(doc, size="A4"):
    """تحديد حجم الصفحة — "A4" أو "Letter" """
    section = doc.sections[0]
    if size == "A4":
        section.page_width  = Inches(8.27)
        section.page_height = Inches(11.69)
    elif size == "Letter":
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
    return doc


def setup_margins(doc, top=1, bottom=1, left=1.25, right=1.25):
    """تحديد هوامش الصفحة بالإنش"""
    section = doc.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)
    return doc


def setup_orientation(doc, orientation="portrait"):
    """اتجاه الصفحة — "portrait" عمودي أو "landscape" أفقي"""
    section = doc.sections[0]
    if orientation == "landscape":
        w = section.page_width
        h = section.page_height
        section.page_width  = h
        section.page_height = w
    return doc


def setup_page_color(doc, hex_color="FFFFFF"):
    """تلوين خلفية الصفحة — مثال: "F0F4FF" """
    bg = doc.element.find(qn('w:background'))
    if bg is None:
        bg = OxmlElement('w:background')
        doc.element.insert(0, bg)
    bg.set(qn('w:color'), hex_color)
    settings = doc.settings.element
    disp_bg = OxmlElement('w:displayBackgroundShape')
    settings.append(disp_bg)
    return doc


def setup_page_border(doc, color="2E75B6", size=6, style="single"):
    """إضافة إطار حول الصفحة — color بصيغة hex، size السماكة"""
    section  = doc.sections[0]
    sectPr   = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   style)
        el.set(qn('w:sz'),    str(size * 8))
        el.set(qn('w:space'), '24')
        el.set(qn('w:color'), color)
        pgBorders.append(el)
    sectPr.append(pgBorders)
    return doc


# ════════════════════════════════════════════════════════════
#  2. الترويسة والتذييل
# ════════════════════════════════════════════════════════════

def add_header(doc, text="", logo_path=None):
    """إضافة ترويسة — نص اختياري وشعار اختياري"""
    section = doc.sections[0]
    header  = section.header
    para    = header.paragraphs[0]
    if logo_path:
        run = para.add_run()
        run.add_picture(logo_path, width=Inches(1.0))
        para.add_run("  ")
    if text:
        run = para.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc


def add_footer(doc, text="سري وخاص"):
    """إضافة تذييل في أسفل كل صفحة"""
    section = doc.sections[0]
    footer  = section.footer
    para    = footer.paragraphs[0]
    run = para.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc


def add_page_number(doc):
    """إضافة رقم الصفحة في التذييل — صيغة: صفحة 1 من 5"""
    section = doc.sections[0]
    footer  = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("صفحة ")
    run.font.size = Pt(9)
    _add_field(para, "PAGE")
    run2 = para.add_run(" من ")
    run2.font.size = Pt(9)
    _add_field(para, "NUMPAGES")
    return doc


def _add_field(para, field_type):
    """دالة مساعدة: تضيف حقل Word مثل PAGE أو NUMPAGES"""
    r1 = para.add_run()
    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), 'begin')
    r1._r.append(fc)

    r2 = para.add_run()
    it = OxmlElement('w:instrText')
    it.text = f' {field_type} '
    r2._r.append(it)

    r3 = para.add_run()
    fe = OxmlElement('w:fldChar')
    fe.set(qn('w:fldCharType'), 'end')
    r3._r.append(fe)


# ════════════════════════════════════════════════════════════
#  3. صفحة الغلاف
# ════════════════════════════════════════════════════════════

def add_cover_logo(doc, logo_path, width=2.0):
    """إضافة شعار الشركة في الغلاف"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(logo_path, width=Inches(width))
    _spacing(para, before=0, after=12)
    return doc


def add_cover_title(doc, text):
    """إضافة العنوان الرئيسي الكبير"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    _spacing(para, before=200, after=12)
    return doc


def add_cover_subtitle(doc, text):
    """إضافة عنوان فرعي تحت العنوان الرئيسي"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size   = Pt(18)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    _spacing(para, before=6, after=6)
    return doc


def add_cover_author(doc, name):
    """إضافة اسم المؤلف"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"إعداد: {name}")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    _spacing(para, before=40, after=4)
    return doc


def add_cover_date(doc, date):
    """إضافة التاريخ — مثال: "يناير 2025" """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(date)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _spacing(para, before=4, after=4)
    return doc


def add_cover_department(doc, name):
    """إضافة اسم الجهة أو القسم"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(name)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    _spacing(para, before=4, after=100)
    return doc


def _spacing(para, before=0, after=0):
    """دالة مساعدة: ضبط المسافة قبل الفقرة وبعدها"""
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before * 20))
    sp.set(qn('w:after'),  str(after  * 20))
    pPr.append(sp)


# ════════════════════════════════════════════════════════════
#  4. المحتوى الأساسي
# ════════════════════════════════════════════════════════════

def add_toc(doc, title="فهرس المحتويات"):
    """إضافة فهرس تلقائي — اضغط "تحديث الحقول" في Word لعرضه"""
    doc.add_paragraph(title).style = "Heading 1"
    para = doc.add_paragraph()
    run  = para.add_run()
    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), 'begin')
    run._r.append(fc)
    it = OxmlElement('w:instrText')
    it.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(it)
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'end')
    run._r.append(fc2)
    return doc


def add_page_break(doc):
    """إضافة فاصل صفحة"""
    doc.add_paragraph().add_run().add_break()
    return doc


def add_heading1(doc, text):
    """عنوان رئيسي H1"""
    para = doc.add_heading(text, level=1)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return doc


def add_heading2(doc, text):
    """عنوان فرعي H2"""
    para = doc.add_heading(text, level=2)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return doc


def add_heading3(doc, text):
    """عنوان فرعي H3"""
    para = doc.add_heading(text, level=3)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return doc


def add_paragraph(doc, text, align="right"):
    """فقرة نص عادي — align: right / left / center / justify"""
    alignments = {
        "right":   WD_ALIGN_PARAGRAPH.RIGHT,
        "left":    WD_ALIGN_PARAGRAPH.LEFT,
        "center":  WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    para = doc.add_paragraph(text)
    para.alignment = alignments.get(align, WD_ALIGN_PARAGRAPH.RIGHT)
    para.runs[0].font.size = Pt(12)
    return doc


def add_rich_text(doc, text, bold=False, italic=False, underline=False, size=12):
    """نص بتنسيق مخصص — غامق / مائل / مسطر"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(text)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.underline = underline
    run.font.size      = Pt(size)
    return doc


def add_colored_text(doc, text, color="2E75B6", highlight=None):
    """نص ملوّن أو مظلّل — highlight: YELLOW / GREEN / RED / BLUE ..."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(text)
    run.font.size = Pt(12)
    r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    if highlight:
        hl_map = {
            "YELLOW":    WD_COLOR_INDEX.YELLOW,
            "GREEN":     WD_COLOR_INDEX.BRIGHT_GREEN,
            "CYAN":      WD_COLOR_INDEX.TURQUOISE,
            "MAGENTA":   WD_COLOR_INDEX.PINK,
            "BLUE":      WD_COLOR_INDEX.BLUE,
            "RED":       WD_COLOR_INDEX.RED,
            "DARK_BLUE": WD_COLOR_INDEX.DARK_BLUE,
            "TURQUOISE": WD_COLOR_INDEX.TURQUOISE,
        }
        run.font.highlight_color = hl_map.get(highlight.upper())
    return doc


def add_bullet_list(doc, items):
    """قائمة نقطية — items: قائمة نصوص"""
    for item in items:
        para = doc.add_paragraph(item, style="List Bullet")
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return doc


def add_numbered_list(doc, items):
    """قائمة مرقّمة — items: قائمة نصوص"""
    for item in items:
        para = doc.add_paragraph(item, style="List Number")
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return doc


def add_nested_list(doc, items):
    """قائمة متداخلة — items: [{"text": "...", "level": 0}, ...]"""
    for item in items:
        level = item.get("level", 0)
        style = "List Bullet 2" if level > 0 else "List Bullet"
        para  = doc.add_paragraph(item["text"], style=style)
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr   = para._p.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            ilvl = numPr.find(qn('w:ilvl'))
            if ilvl is not None:
                ilvl.set(qn('w:val'), str(level))
    return doc


# ════════════════════════════════════════════════════════════
#  5. الجداول
# ════════════════════════════════════════════════════════════

def create_table(doc, data, col_widths=None):
    """
    إنشاء جدول من بيانات ثنائية الأبعاد.
    data = [["رأس1","رأس2"], ["قيمة1","قيمة2"]]
    col_widths = [2.5, 2.5]  بالإنش (اختياري)
    يُرجع (doc, table)
    """
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = "Table Grid"
    for r, row_data in enumerate(data):
        for c, cell_text in enumerate(row_data):
            cell = table.cell(r, c)
            cell.text = str(cell_text)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = Inches(col_widths[i])
    return doc, table


def style_table_header(table, row_index=0, bg_color="1F3964", text_color="FFFFFF"):
    """تلوين صف رأس الجدول"""
    for cell in table.rows[row_index].cells:
        _set_cell_bg(cell, bg_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                r = int(text_color[0:2], 16)
                g = int(text_color[2:4], 16)
                b = int(text_color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
    return table


def set_cell_color(table, row, col, hex_color):
    """تلوين خلية واحدة"""
    _set_cell_bg(table.cell(row, col), hex_color)
    return table


def set_table_borders(table, color="2E75B6", size=4):
    """تخصيص لون وسماكة حدود الجدول كاملاً"""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'),   'single')
                border.set(qn('w:sz'),    str(size))
                border.set(qn('w:color'), color)
                tcBorders.append(border)
            tcPr.append(tcBorders)
    return table


def set_cell_alignment(table, row, col, h_align="center", v_align="center"):
    """ضبط محاذاة خلية — h: left/center/right | v: top/center/bottom"""
    cell = table.cell(row, col)
    h_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}
    v_map = {"top": WD_ALIGN_VERTICAL.TOP, "center": WD_ALIGN_VERTICAL.CENTER, "bottom": WD_ALIGN_VERTICAL.BOTTOM}
    for para in cell.paragraphs:
        para.alignment = h_map.get(h_align, WD_ALIGN_PARAGRAPH.CENTER)
    cell.vertical_alignment = v_map.get(v_align, WD_ALIGN_VERTICAL.CENTER)
    return table


def merge_cells(table, r1, c1, r2, c2):
    """دمج خلايا من (r1,c1) إلى (r2,c2)"""
    table.cell(r1, c1).merge(table.cell(r2, c2))
    return table


def _set_cell_bg(cell, hex_color):
    """دالة مساعدة: تلوين خلفية خلية"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


# ════════════════════════════════════════════════════════════
#  6. الصور والوسائط
# ════════════════════════════════════════════════════════════

def add_image(doc, image_path, width=4.0, align="center"):
    """إدراج صورة — width بالإنش، align: left/center/right"""
    alignments = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}
    para = doc.add_paragraph()
    para.alignment = alignments.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    para.add_run().add_picture(image_path, width=Inches(width))
    return doc


def add_image_caption(doc, text):
    """تسمية توضيحية تحت الصورة"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"شكل: {text}")
    run.font.size   = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return doc


def add_text_box(doc, text, bg_color="EBF3FB", border_color="2E75B6"):
    """مربع نص ملوّن (يُنفَّذ كجدول خلية واحدة)"""
    table = doc.add_table(rows=1, cols=1)
    cell  = table.cell(0, 0)
    cell.text = text
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in para.runs:
            run.font.size = Pt(11)
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  bg_color)
    tcPr.append(shd)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '6')
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)
    return doc


# ════════════════════════════════════════════════════════════
#  7. العناصر المتقدمة
# ════════════════════════════════════════════════════════════

def add_section_break(doc, break_type="next_page"):
    """فاصل قسم — "next_page" أو "continuous" """
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    sectPr  = OxmlElement('w:sectPr')
    type_el = OxmlElement('w:type')
    type_el.set(qn('w:val'), "nextPage" if break_type == "next_page" else "continuous")
    sectPr.append(type_el)
    pPr.append(sectPr)
    return doc


def add_footnote(doc, paragraph, footnote_text):
    """حاشية سفلية مبسطة — تضيف * في النص ونصاً في الأسفل"""
    run = paragraph.add_run("*")
    run.font.superscript = True
    run.font.size = Pt(8)
    note_run = doc.add_paragraph().add_run(f"* {footnote_text}")
    note_run.font.size   = Pt(9)
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return doc


def add_comment(doc, paragraph, comment_text):
    """تعليق مرئي على فقرة (مبسط)"""
    run = paragraph.add_run(f" [تعليق: {comment_text}]")
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.font.size      = Pt(10)
    run.font.italic    = True
    return doc


def add_hyperlink(doc, text, url):
    """رابط تشعبي قابل للنقر"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_id = para.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run_el = OxmlElement('w:r')
    rPr    = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    run_el.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run_el.append(t)
    hyperlink.append(run_el)
    para._p.append(hyperlink)
    return doc


def add_horizontal_rule(doc, color="2E75B6", size=6):
    """خط فاصل أفقي"""
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return doc



#===============
#compiling all 
#===============

def compile_report_from_ai_json(json_string, output_filename):
    try:
        json_string = str(json_string).strip()
        
        if json_string.startswith("```json"):
            json_string = json_string[7:]
        if json_string.endswith("```"):
            json_string = json_string[:-3]
            
        json_string = json_string.strip()

        data = json.loads(json_string) 
    except json.JSONDecodeError as e :
        print("error : {e} \n please try again...")
        return
    
    doc = Document()

    # 1. Page Setup
    ps = data.get("page_setup", {})
    doc = setup_page_size(doc, ps.get("size", "A4"))
    doc = setup_margins(doc, top=ps.get("top_margin", 1), bottom=ps.get("bottom_margin", 1), left=ps.get("left_margin", 1.25), right=ps.get("right_margin", 1.25))
    doc = setup_orientation(doc, ps.get("orientation", "portrait"))
    doc = setup_page_color(doc, ps.get("page_color", "FFFFFF"))

    if "border" in ps and ps["border"] is not None:
            b = ps["border"]
            if isinstance(b, dict):
                doc = setup_page_border(
                    doc, 
                    color=b.get("color", "1F3964"), 
                    size=b.get("size", 4), 
                    style=b.get("style", "single")
                )

    # 2. Header & Footer
    hf = data.get("header_footer", {})
    if "header" in hf:
        doc = add_header(doc, text=hf["header"].get("text", "")) #logo_path=hf["header"].get("logo_path")
    if "footer" in hf:
        doc = add_footer(doc, text=hf["footer"].get("text", "Confidential"))
    if hf.get("include_page_number", True):
        doc = add_page_number(doc)

    # 3. Cover Page
    cp = data.get("cover_page", {})
    if cp.get("has_cover", False):
       # if cp.get("logo_path"):
        #    doc = add_cover_logo(doc, logo_path=cp["logo_path"], width=cp.get("logo_width", 2.0))
        doc = add_cover_title(doc, cp.get("title", "Report"))
        if cp.get("subtitle"): doc = add_cover_subtitle(doc, cp["subtitle"])
        if cp.get("author"): doc = add_cover_author(doc, cp["author"])
        if cp.get("date"): doc = add_cover_date(doc, cp["date"])
        if cp.get("department"): doc = add_cover_department(doc, cp["department"])

    # 4. Sequential Elements Loop
    for element in data.get("report_structure", []):
        el_type = element.get("type")
        el_data = element.get("data")
        
        if el_type == "toc":
            doc = add_toc(doc, title=el_data.get("title", "Table of Contents"))
        elif el_type == "page_break":
            doc = add_page_break(doc)
        elif el_type == "heading1":
            doc = add_heading1(doc, el_data)
        elif el_type == "heading2":
            doc = add_heading2(doc, el_data)
        elif el_type == "heading3":
            doc = add_heading3(doc, el_data)
        elif el_type == "paragraph":
            doc = add_paragraph(doc, el_data.get("text"), align=el_data.get("align", "right"))
        elif el_type == "rich_text":
            doc = add_rich_text(doc, el_data.get("text"), bold=el_data.get("bold", False), italic=el_data.get("italic", False), underline=el_data.get("underline", False), size=el_data.get("size", 12))
        elif el_type == "colored_text":
            doc = add_colored_text(doc, el_data.get("text"), color=el_data.get("color", "2E75B6"), highlight=el_data.get("highlight"))
        elif el_type == "text_box":
            doc = add_text_box(doc, el_data.get("text"), bg_color=el_data.get("bg_color", "EBF3FB"), border_color=el_data.get("border_color", "2E75B6"))
        elif el_type == "bullet_list":
            doc = add_bullet_list(doc, el_data)
        elif el_type == "numbered_list":
            doc = add_numbered_list(doc, el_data)
        elif el_type == "nested_list":
            doc = add_nested_list(doc, el_data)
        elif el_type == "horizontal_rule":
            doc = add_horizontal_rule(doc, color=el_data.get("color", "2E75B6"), size=el_data.get("size", 6))
        elif el_type == "hyperlink":
            doc = add_hyperlink(doc, el_data.get("text"), el_data.get("url"))
        #elif el_type == "image":
         #   doc = add_image(doc, image_path=el_data.get("image_path"), width=el_data.get("width", 4.0), align=el_data.get("align", "center"))
          #  if el_data.get("caption"):
           #     doc = add_image_caption(doc, el_data["caption"])
        elif el_type == "section_break":
            doc = add_section_break(doc, break_type=el_data.get("break_type", "next_page"))
            
        # معالجة الجداول المعقدة مالتك بالتفصيل والأوامر الفرعية داخل الـ JSON
        elif el_type == "table":
            doc, table = create_table(doc, el_data.get("matrix"), col_widths=el_data.get("col_widths"))
            if "header_style" in el_data:
                hs = el_data["header_style"]
                table = style_table_header(table, row_index=0, bg_color=hs.get("bg_color", "1F3964"), text_color=hs.get("text_color", "FFFFFF"))
            if "borders" in el_data:
                tb = el_data["borders"]
                table = set_table_borders(table, color=tb.get("color", "2E75B6"), size=tb.get("size", 4))
            for cell in el_data.get("cell_customization", []):
                if "bg_color" in cell:
                    table = set_cell_color(table, row=cell["row"], col=cell["col"], hex_color=cell["bg_color"])
                table = set_cell_alignment(table, row=cell["row"], col=cell["col"], h_align=cell.get("h_align", "center"), v_align=cell.get("v_align", "center"))
            for merge in el_data.get("merges", []):
                table = merge_cells(table, r1=merge["r1"], c1=merge["c1"], r2=merge["r2"], c2=merge["c2"])
    doc.save(output_filename)
    return output_filename


# ===============
# making excel 
#================


def create_excel_report_from_json(json_string):
    try:
        # ── Step 1: Clean JSON from AI ───────────────────
        json_string = str(json_string).strip()
        if json_string.startswith("```json"):
            json_string = json_string[7:]
        if json_string.endswith("```"):
            json_string = json_string[:-3]
        json_string = json_string.strip()
        
        try:
            data = json.loads(json_string)
        except json.JSONDecodeError as e:
            print(f"❌ Failed to parse Excel JSON from AI: {e}")
            return None

        # 1. Create workbook and active sheet
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = data.get("sheet_title", "Sheet1")
        
        # ← حذفنا السطر RTL من هنا

        # 2. Fonts & Fills
        theme_color = data.get("theme_color")
        if not theme_color:
            theme_color = "1F3964"
        theme_color = theme_color.lstrip('#')
        
        header_fill  = PatternFill(start_color=theme_color, end_color=theme_color, fill_type="solid")
        header_font  = Font(name="Arial", size=11, bold=True, color="FFFFFF")
        data_font    = Font(name="Arial", size=11)
        center_align = Alignment(horizontal="center", vertical="center")
        
        thin_border = Border(
            left=Side(style='thin',   color='D9D9D9'),
            right=Side(style='thin',  color='D9D9D9'),
            top=Side(style='thin',    color='D9D9D9'),
            bottom=Side(style='thin', color='D9D9D9')
        )

        # 3. Write headers
        headers = data.get("headers", [])
        if headers:
            ws.append(headers)
            for cell in ws[1]:
                cell.font      = header_font
                cell.fill      = header_fill
                cell.alignment = center_align
                cell.border    = thin_border
            ws.row_dimensions[1].height = 25

        # 4. Write rows
        rows = data.get("rows", [])
        for row in rows:
            ws.append(row)
            
        # 5. Format data cells
        if ws.max_row >= 2:
            for row_idx, row in enumerate(ws.iter_rows(min_row=2, max_row=ws.max_row), start=2):
                ws.row_dimensions[row_idx].height = 20
                for cell in row:
                    cell.font      = data_font
                    cell.alignment = center_align
                    cell.border    = thin_border
                    
            # Zebra striping
            for row_idx in range(2, ws.max_row + 1, 2):
                for cell in ws[row_idx]:
                    cell.fill = PatternFill(start_color="F2F5F9", end_color="F2F5F9", fill_type="solid")

        # 6. Auto-fit columns
        for col in ws.columns:
            max_len = 0
            for cell in col:
                val_str = str(cell.value or '')
                if val_str.startswith('='):
                    val_str = "12345678"
                if len(val_str) > max_len:
                    max_len = len(val_str)
            col_letter = get_column_letter(col[0].column)
            ws.column_dimensions[col_letter].width = max(max_len + 4, 12)

        # 7. Save to Downloads
        filename = data.get("filename", "output.xlsx")
        downloads_path = os.path.join(os.path.expanduser('~'), 'Downloads', filename)
        
        try:
            wb.save(downloads_path)
            print(f"Excel file generated successfully: {downloads_path}")
            return downloads_path
        except IOError:
            print(f" Error: {filename} is currently open! Please close it first.")
            backup_path = downloads_path.replace(".xlsx", "_backup.xlsx")
            wb.save(backup_path)
            return backup_path
        
    except Exception as e :
        print(f"error : {e}")


#=======================
#converting docx to pdf
#=======================


def convert_word_to_pdf_cross_platform(word_file):
    try:
        file_path = controltools.get_file_path_hybrid(word_file)
        if not file_path:
            print("File not found.")
            return None

        output_dir = os.path.dirname(os.path.abspath(file_path))
        print("Converting your file...")

        if platform.system() == "Windows":
            # يحتاج Microsoft Word مثبت
            convert(file_path)
        else:
            # يحتاج LibreOffice مثبت
            subprocess.run([
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", output_dir,
                file_path
            ], check=True)

        pdf_path = os.path.join(output_dir, os.path.basename(file_path).replace('.docx', '.pdf'))
        print(f"Done! File path: {pdf_path}")
        return pdf_path

    except Exception as e:
        print(f"An error occurred: {e}")
        return None
    

def make_excel_from_file(target, query):
    file_path = controltools.get_file_path_hybrid(target)
    if not file_path:
        print("File not found.")
        return
    
    file_content = controltools.read_file_content(file_path)
    if not file_content:
        print("Unsupported file type.")
        return
    
    print("Reading file and generating Excel, please wait...")
    
    excel_request_prompt = f"""
    {excel_prompt}
    
    USER REQUEST: {query}
    
    HERE IS THE ACTUAL DATA FROM THE FILE, USE IT EXACTLY:
    {file_content}
    """
    
    excel_response = client.models.generate_content(
        model='gemini-2.5-flash',
        contents=excel_request_prompt
    )
    
    create_excel_report_from_json(excel_response.text)